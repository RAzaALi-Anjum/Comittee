import docx
import sys

doc = docx.Document('Test Cases.docx')

output = []

# Extract paragraphs
for para in doc.paragraphs:
    if para.text.strip():
        output.append(para.text)

# Extract tables
for i, table in enumerate(doc.tables):
    output.append(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip().replace('\n', ' | '))
        output.append('\t|\t'.join(row_data))

full_text = '\n'.join(output)

# Write with utf-8 encoding
with open('extracted_tests.txt', 'w', encoding='utf-8') as f:
    f.write(full_text)

print(f"Done. Total chars: {len(full_text)}")
